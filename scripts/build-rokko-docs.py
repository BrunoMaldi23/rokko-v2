import json
import re
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def clean_text(value: str) -> str:
    value = value or ""
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def safe_name(value: str, fallback: str = "Sin asunto") -> str:
    value = re.sub(r'[<>:"/\\|?*]+', "_", value or fallback)
    value = re.sub(r"\s+", " ", value).strip()
    return (value or fallback)[:80].strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def apply_base_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def add_title(doc: Document, title: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(14)
        sr = sp.add_run(subtitle)
        sr.font.name = "Calibri"
        sr.font.size = Pt(10)
        sr.font.color.rgb = RGBColor.from_string("555555")


def add_meta_table(doc: Document, rows) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for label, value in rows:
        row = table.add_row()
        row.cells[0].width = Inches(1.55)
        row.cells[1].width = Inches(4.95)
        row.cells[0].text = label
        row.cells[1].text = value or ""
        set_cell_shading(row.cells[0], "F2F4F7")
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def add_email_content(doc: Document, message: dict, include_break: bool = False) -> None:
    if include_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading(message.get("subject") or "Sin asunto", level=1)
    received = message.get("received_time", "")
    try:
        received = datetime.fromisoformat(received).strftime("%Y-%m-%d %H:%M")
    except Exception:
        pass

    add_meta_table(
        doc,
        [
            ("Remitente", f"{message.get('sender_name', '')} <{message.get('sender_email', '')}>"),
            ("Fecha", received),
            ("Carpeta", message.get("folder", "")),
            ("Categoria", message.get("category", "General")),
            ("Imagenes guardadas", str(len(message.get("attachments", [])))),
        ],
    )

    body = clean_text(message.get("body", ""))
    if body:
        doc.add_heading("Contenido", level=2)
        for block in body.split("\n\n"):
            paragraph = doc.add_paragraph()
            paragraph.add_run(block.strip())
    else:
        doc.add_paragraph("Este correo no tiene contenido de texto legible en Outlook.")

    attachments = message.get("attachments", [])
    if attachments:
        doc.add_heading("Imagenes adjuntas guardadas", level=2)
        for attachment in attachments:
            p = doc.add_paragraph(style=None)
            p.style = doc.styles["List Bullet"]
            p.add_run(f"{attachment.get('file_name', '')} - {attachment.get('category', '')}")


def save_doc(messages, output_path: Path, title: str, subtitle: str) -> None:
    doc = Document()
    apply_base_styles(doc)
    add_title(doc, title, subtitle)

    if not messages:
        doc.add_paragraph("No se encontraron correos para este documento.")
    else:
        for idx, message in enumerate(messages):
            add_email_content(doc, message, include_break=idx > 0)

    doc.save(output_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: build-rokko-docs.py export.json output_root", file=sys.stderr)
        return 2

    export_path = Path(sys.argv[1])
    output_root = Path(sys.argv[2])
    docs_root = output_root / "documentos"
    docs_root.mkdir(parents=True, exist_ok=True)
    category_docs_root = docs_root / "por categoria"
    category_docs_root.mkdir(parents=True, exist_ok=True)

    data = json.loads(export_path.read_text(encoding="utf-8-sig"))
    if isinstance(data, dict):
        messages = [data]
    else:
        messages = data

    now_label = datetime.now().strftime("%Y-%m-%d %H:%M")
    subtitle = f"Exportado desde Outlook el {now_label}. Remitente filtrado: contacto.rokko@gmail.com"
    consolidated = docs_root / "Resumen correos Rokko.docx"
    save_doc(messages, consolidated, "Resumen correos Rokko", subtitle)

    grouped = defaultdict(list)
    for message in messages:
        grouped[message.get("category") or "General"].append(message)

    category_paths = []
    for category, category_messages in sorted(grouped.items()):
        path = category_docs_root / f"{safe_name(category)}.docx"
        save_doc(category_messages, path, f"Correos Rokko - {category}", subtitle)
        category_paths.append(str(path))

    print(json.dumps({
        "consolidated_doc": str(consolidated),
        "category_docs": category_paths,
        "message_count": len(messages),
        "category_count": len(grouped),
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
